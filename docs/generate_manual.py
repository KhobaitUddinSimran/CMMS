"""
Generate CMMS User Manual as a .docx file.
Run: python3 docs/generate_manual.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

UTM_RED = RGBColor(0xC9, 0x00, 0x31)
DARK    = RGBColor(0x11, 0x18, 0x27)
GREY    = RGBColor(0x6B, 0x72, 0x80)
LIGHT   = RGBColor(0xF9, 0xFA, 0xFB)


def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=UTM_RED)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=DARK)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=DARK)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def note_box(doc, text, label="Note"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True, size=10.5, color=UTM_RED)
    r2 = p.add_run(text)
    set_font(r2, size=10.5, color=GREY, italic=True)


def step_table(doc, steps):
    """Render numbered steps in a borderless table."""
    t = doc.add_table(rows=len(steps), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, step in enumerate(steps):
        t.rows[i].cells[0].width = Inches(0.45)
        t.rows[i].cells[1].width = Inches(5.5)
        n_para = t.rows[i].cells[0].paragraphs[0]
        nr = n_para.add_run(f"{i+1}.")
        set_font(nr, bold=True, size=11, color=UTM_RED)
        s_para = t.rows[i].cells[1].paragraphs[0]
        sr = s_para.add_run(step)
        set_font(sr, size=11)
    doc.add_paragraph()


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("CMMS")
set_font(r, bold=True, size=42, color=UTM_RED)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Course Mark Management System")
set_font(r2, bold=True, size=22, color=DARK)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("User Manual  ·  Version 5.0  ·  May 2026")
set_font(r3, size=13, color=GREY)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Universiti Teknologi Malaysia")
set_font(r4, size=12, color=GREY, italic=True)

doc.add_page_break()

# ── Table of Contents (manual) ─────────────────────────────────────────────
heading1(doc, "Table of Contents")
toc_items = [
    ("1", "System Overview"),
    ("2", "Getting Started — Login & Account Approval"),
    ("3", "Role Guide: Student"),
    ("4", "Role Guide: Lecturer"),
    ("5", "Role Guide: Coordinator"),
    ("6", "Role Guide: Head of Department (HOD)"),
    ("7", "Role Guide: Administrator"),
    ("8", "Feature Reference"),
    ("9", "Frequently Asked Questions"),
]
for num, label in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"  {num}.  {label}")
    set_font(r, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "1. System Overview")
body(doc, (
    "CMMS (Course Mark Management System) is a web-based application developed for Universiti Teknologi Malaysia "
    "to streamline the management of coursework marks, assessments, student queries, and semester timelines across "
    "all academic departments."
))
body(doc, "The system serves five distinct user roles:")

roles_info = [
    ("Student",      "Views enrolled courses, marks, and grade progress; submits queries about marks."),
    ("Lecturer",     "Enters and publishes marks; configures assessments; responds to student queries; imports marks from Excel."),
    ("Coordinator",  "Manages courses and lecturer assignments; configures semester timelines; oversees flagged marks and reports."),
    ("HOD",          "Reviews departmental analytics, course management, teaching workloads, and audit logs."),
    ("Admin",        "Manages user accounts, approvals, roles, and system configuration."),
]
for role, desc in roles_info:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(f"{role}: ")
    set_font(r1, bold=True, size=11)
    r2 = p.add_run(desc)
    set_font(r2, size=11)

divider(doc)
heading2(doc, "Technology Stack")
bullet(doc, "Frontend: Next.js 14 (React 18, TypeScript 5.3, Tailwind CSS)")
bullet(doc, "Backend: FastAPI (Python) — deployed on Render.com")
bullet(doc, "Database: Supabase PostgreSQL")
bullet(doc, "Authentication: JWT-based with role claims")
body(doc, "Access the application at your institution's deployment URL.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "2. Getting Started — Login & Account Approval")

heading2(doc, "2.1  Signing Up")
body(doc, "New users must register before accessing the system. Only Students and Lecturers can self-register; Coordinators, HODs, and Admins are created by an existing Administrator.")
step_table(doc, [
    "Navigate to the CMMS login page and click Create Account.",
    "Select your role: Student or Lecturer.",
    "Enter your full name, institutional email address, and a secure password "
    "(minimum 8 characters, at least one uppercase letter and one number).",
    "Students must use @graduate.utm.my; Lecturers must use @utm.my.",
    "Submit the form. Your account will be placed in Pending Approval status.",
    "Wait for an Administrator to approve your account. You will be notified on next login.",
])
note_box(doc, "Your account is not active until an Administrator approves it. The system polls for approval status every 30 seconds after login.")

heading2(doc, "2.2  Logging In")
step_table(doc, [
    "Go to the CMMS home page.",
    "Select your role from the role picker.",
    "Enter your email address and password.",
    "Click Sign In.",
    "If approved, you are redirected to your role-specific Dashboard.",
    "If pending or rejected, you will see the corresponding status screen.",
])

heading2(doc, "2.3  Resetting Your Password")
step_table(doc, [
    "On the login page, click Forgot Password.",
    "Enter your registered email address and click Send Reset Link.",
    "Check your email for the reset link (valid for 10 minutes).",
    "Follow the link, enter and confirm your new password, then submit.",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. STUDENT GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "3. Role Guide: Student")

heading2(doc, "3.1  Dashboard")
body(doc, "After login, students land on the Student Dashboard which shows a summary of enrolled courses and any outstanding queries.")

heading2(doc, "3.2  My Courses")
body(doc, "The My Courses page lists all courses the student is enrolled in for the current semester.")
bullet(doc, "Click a course name to view its details.")
bullet(doc, "Each course card shows the course code, section, semester, and assigned lecturer.")

heading2(doc, "3.3  My Marks")
body(doc, "The My Marks page displays the student's published scores for all assessments across all enrolled courses.")
bullet(doc, "Only published marks are visible. Draft marks are not shown until the lecturer publishes them.")
bullet(doc, "The Carry % column shows the running weighted contribution of published marks toward the final grade.")
bullet(doc, "A colour indicator shows performance: green (≥ 70%), amber (50–69%), red (< 50%).")

heading2(doc, "3.4  Queries")
body(doc, "Students can submit questions or concerns about their published marks.")
step_table(doc, [
    "Go to Queries in the sidebar and click New Query.",
    "Select the course and then the assessment you are querying.",
    "Write a clear description of your question or concern.",
    "Click Submit.",
    "You will see the query listed with status OPEN.",
    "When a lecturer responds, the query status changes to RESOLVED and a red dot appears on the Queries sidebar link.",
    "Click the query to expand it and read the staff response.",
    "If unsatisfied, click Reopen Query to escalate.",
])
note_box(doc, "A red dot badge on the Queries sidebar link means you have an unread staff response.")

heading2(doc, "3.5  Filtering & Sorting Queries")
bullet(doc, "Use the All / Open / Resolved filter tabs to narrow the list.")
bullet(doc, "Click the Newest first / Oldest first button to toggle sort order.")
bullet(doc, "Each query card shows the full submission date and time.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. LECTURER GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "4. Role Guide: Lecturer")

heading2(doc, "4.1  Dashboard")
body(doc, "The Lecturer Dashboard displays the courses you are assigned to this semester and quick-access statistics: total students, pending marks, and unread queries.")

heading2(doc, "4.2  My Courses")
body(doc, "Lists all courses assigned to you. Click a course to view its details, enrolled students, and assessment setup.")

heading2(doc, "4.3  Import Students (Roster Upload)")
body(doc, "Upload a CSV or Excel file of students to bulk-enrol them into a course.")
step_table(doc, [
    "Go to Import Students in the sidebar.",
    "Select the target course from the dropdown.",
    "Click Choose File and select a CSV or XLSX file containing student matric numbers or email addresses.",
    "Click Preview to review the import before committing.",
    "Click Confirm Import to finalise enrolment.",
])
note_box(doc, "The upload file should have one student per row. The first row is treated as a header and ignored.")

heading2(doc, "4.4  Mark Entry (Grade Sheet)")
body(doc, "The Mark Entry (Grade Sheet) page is the central place for entering and managing student marks.")

heading3(doc, "4.4.1  Selecting a Course")
body(doc, "Choose a course from the course dropdown at the top of the page. The grade grid will load automatically.")

heading3(doc, "4.4.2  Entering Marks")
bullet(doc, "Each cell in the grid corresponds to one student's score for one assessment.")
bullet(doc, "Type a numeric score in the cell. Cells turn amber when a change has not been saved yet.")
bullet(doc, "A cell turns red if the score exceeds the assessment's maximum score.")
bullet(doc, "Click Save All Changes to persist all unsaved marks to the database.")

heading3(doc, "4.4.3  Publishing Marks")
bullet(doc, "Draft marks are only visible to the lecturer. Students see marks only after they are published.")
bullet(doc, "Click Publish in the column header to publish all draft marks for that assessment.")
bullet(doc, "Click Publish All Drafts in the action bar to publish every draft mark across all assessments.")
bullet(doc, "Published marks are shown in green with a ✓ indicator. They cannot be edited until unpublished.")
bullet(doc, "Click Unpublish in the column header to revert published marks to draft for correction.")

heading3(doc, "4.4.4  Importing Marks from Excel")
body(doc, "You can bulk-import student marks from a prepared Excel file.")
step_table(doc, [
    "Open the Grade Sheet for the target course.",
    "Click the Import Excel button in the action bar.",
    "In the dialog, review the expected file format: Column 1 must be student email or matric number; "
    "remaining columns must exactly match your assessment names.",
    "Click Choose File (.xlsx) and select your Excel file.",
    "Click Import.",
    "The system will show how many rows were successfully imported, how many were skipped (no matching student), "
    "and any row-level errors (e.g., invalid score format, unknown assessment name).",
    "Close the dialog. The grid reloads automatically with the imported marks.",
])
note_box(doc, "Imported marks are saved as draft status. You must publish them manually.")

heading3(doc, "4.4.5  Statistics Footer")
body(doc, "At the bottom of the grid, per-assessment statistics are shown: Average, Highest, and Lowest scores across all entered marks (draft + published).")

heading3(doc, "4.4.6  Carry % Column")
body(doc, (
    "The rightmost Carry % column shows the running weighted carry total for each student based on "
    "published marks only. This is calculated as the sum of (score / max_score × weight) for each published mark."
))

heading2(doc, "4.5  Assessment Setup")
body(doc, "Configure the assessment components (quizzes, assignments, midterm, etc.) for each course.")

heading3(doc, "4.5.1  Adding an Assessment")
step_table(doc, [
    "Go to Assessment Setup and select the course.",
    "Ensure the Assessment Config tab is active.",
    "Fill in the assessment name, type, maximum score, and weight percentage.",
    "Click Add Assessment.",
    "Repeat for each assessment component. Ensure total weights sum to 100%.",
])

heading3(doc, "4.5.2  Locking the Schema")
body(doc, "Once all assessments are configured, click Lock Schema. Locking prevents further changes to the assessment structure, ensuring data integrity once marks are being entered.")
note_box(doc, "Locking is irreversible from the UI. Contact a Coordinator or Admin to unlock if necessary.")

heading3(doc, "4.5.3  Carry Marks Summary Tab")
body(doc, "Switch to the Carry Marks Summary tab to view a ranked table of all students' cumulative carry percentages based on published marks.")
bullet(doc, "Students are sorted from highest to lowest carry %.")
bullet(doc, "Grade bands: Strong (≥ 70%), Satisfactory (50–69%), At Risk (< 50%).")
bullet(doc, "Click Refresh to reload the latest data.")

heading2(doc, "4.6  Queries")
body(doc, "Students submit queries about their marks. As a lecturer, you receive and respond to these.")
bullet(doc, "A red badge on the Queries sidebar link shows the number of unread/unanswered queries.")
bullet(doc, "Open a query card to read the student's question.")
bullet(doc, "Click Respond to type and send a response. The query is automatically marked as Resolved.")
bullet(doc, "Click Reopen to change a resolved query back to Open if further discussion is needed.")
bullet(doc, "Use the All / Open / Resolved tabs and the Newest/Oldest sort to manage your queue.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. COORDINATOR GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "5. Role Guide: Coordinator")

heading2(doc, "5.1  Courses")
body(doc, "Coordinators can view and manage all courses in their department.")
bullet(doc, "Click Create Course to add a new course (code, section, semester, credits, maximum students).")
bullet(doc, "Click a course to view its detail page, edit information, or assign a lecturer.")

heading2(doc, "5.2  Course Management — Assigning Lecturers")
body(doc, "Use the Course Management page to assign lecturers to courses.")
step_table(doc, [
    "Open Course Management and select a course.",
    "In the Assign Lecturer section, choose a lecturer from the dropdown.",
    "The lecturer's current credit workload is displayed: credits used / maximum allowed.",
    "If a lecturer has no credit limit set, No limit is shown.",
    "If the assignment would exceed the credit limit, a warning appears. Coordinators may override with a reason.",
    "Click Assign to confirm.",
])
note_box(doc, "Per-lecturer credit limits are configurable by Administrators in the user profile settings.")

heading2(doc, "5.3  Roster Management")
body(doc, "Manage student enrolments for courses.")
bullet(doc, "Upload bulk roster files (CSV/XLSX).")
bullet(doc, "Preview before confirming to avoid errors.")
bullet(doc, "Drop individual students if needed.")

heading2(doc, "5.4  Semester Timeline")
body(doc, "Set key academic deadlines for each semester.")
step_table(doc, [
    "Go to Semester Timeline in the sidebar.",
    "Click the + New Timeline button or click an existing timeline to edit it.",
    "Set the Academic Year, Semester, Start Date, End Date, Midterm Deadline, and Grade Submission Deadline.",
    "Optionally add notes.",
    "Click Save.",
])
body(doc, "Timelines are displayed as colour-coded horizontal bars on the calendar. Upcoming deadlines are highlighted with badges.")
note_box(doc, "Finals deadlines have been removed. Only coursework deadlines (midterm and grade submission) are tracked.")

heading2(doc, "5.5  Flagged Marks")
body(doc, "Marks flagged by the system (e.g., anomalies detected by the AI flag) or by lecturers appear here.")
bullet(doc, "Review the flag note for context.")
bullet(doc, "Take action: clear the flag, contact the lecturer, or escalate.")

heading2(doc, "5.6  Reports")
body(doc, "Generate grade distribution reports and export data for departmental review.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. HOD GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "6. Role Guide: Head of Department (HOD)")

heading2(doc, "6.1  Dashboard")
body(doc, "The HOD Dashboard provides a high-level overview of department performance, active courses, and pending actions.")

heading2(doc, "6.2  Course Management")
body(doc, "HODs have full access to course management, including assigning lecturers and overriding credit limits.")
bullet(doc, "View all courses and their assigned lecturers.")
bullet(doc, "Assign or reassign lecturers with the same workflow as Coordinators.")
bullet(doc, "Override the credit limit for a lecturer by providing a justification reason.")

heading2(doc, "6.3  Analytics")
body(doc, "Departmental analytics show grade distributions, at-risk students, and overall course performance.")

heading2(doc, "6.4  Export")
body(doc, "Export course data, grade summaries, and roster information to CSV or Excel for external reporting.")

heading2(doc, "6.5  Flagged Marks")
body(doc, "Review and take action on all flagged marks across the department.")

heading2(doc, "6.6  Audit Log")
body(doc, "The Audit Log provides a complete record of all significant system actions (mark updates, schema locks, user approvals, etc.) for accountability and compliance.")
bullet(doc, "Filter by date, user, or action type.")
bullet(doc, "Each entry shows the before/after values of changed records.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. ADMIN GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "7. Role Guide: Administrator")

heading2(doc, "7.1  Pending Approvals")
body(doc, "New self-registered users (Students and Lecturers) must be approved before they can access the system.")
step_table(doc, [
    "Go to Pending Approvals.",
    "Review the applicant's details (name, email, role).",
    "Click Approve to activate the account, or Reject (with a reason) to deny access.",
])

heading2(doc, "7.2  Users")
body(doc, "View and manage all system users.")
bullet(doc, "Search by name or email.")
bullet(doc, "View user details, roles, and account status.")
bullet(doc, "Deactivate or reactivate accounts.")
bullet(doc, "Set per-lecturer maximum teaching credit limits from the user profile page.")

heading2(doc, "7.3  Roles & Permissions")
body(doc, "Assign or revoke special roles (Coordinator, HOD) to existing Lecturer accounts.")
step_table(doc, [
    "Go to Roles & Permissions.",
    "Find the lecturer's account.",
    "Click Assign Special Role and select Coordinator or HOD.",
    "To remove a role, click Revoke Role.",
])
note_box(doc, "Only one special role can be active at a time per user in the current implementation.")

heading2(doc, "7.4  Per-Lecturer Credit Limit")
body(doc, "Administrators can set a flexible maximum teaching credit limit for each lecturer.")
step_table(doc, [
    "Go to Users and open the lecturer's profile.",
    "Find the Max Teaching Credits field.",
    "Enter the maximum number of credits the lecturer may be assigned per semester, or leave blank for no limit.",
    "Save the profile.",
])
note_box(doc, "If no limit is set, the lecturer selector shows 'No limit' and the system will not block assignments based on credits.")

heading2(doc, "7.5  System Settings & Logs")
body(doc, "Configure system-wide settings and review system logs for debugging and auditing.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. FEATURE REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "8. Feature Reference")

heading2(doc, "8.1  Excel Import for Marks")
body(doc, "Lecturers can import student marks in bulk from a Microsoft Excel (.xlsx) file.")
heading3(doc, "File Format Requirements")
bullet(doc, "Row 1: Header row. Column A header must be email or matric_number. Remaining columns must match assessment names exactly (case-insensitive).")
bullet(doc, "Row 2 onwards: One student per row. Column A contains the student identifier; other columns contain numeric scores.")
bullet(doc, "Scores must be numeric. Empty cells are skipped (no mark created).")
heading3(doc, "Import Behaviour")
bullet(doc, "If a mark already exists for the student/assessment pair, it is updated (upserted).")
bullet(doc, "If the student is not enrolled in the course, that row is skipped with an error message.")
bullet(doc, "If an assessment name does not match any configured assessment, that column is skipped.")
bullet(doc, "Imported marks have draft status and must be published manually.")
heading3(doc, "Example File Structure")
t = doc.add_table(rows=3, cols=4)
t.style = 'Table Grid'
headers = ["email", "Quiz 1", "Assignment 1", "Midterm"]
rows_data = [
    ["alice@graduate.utm.my", "18", "85", "72"],
    ["bob@graduate.utm.my",   "15", "78", ""],
]
for j, h in enumerate(headers):
    cell = t.rows[0].cells[j]
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(h)
    set_font(r, bold=True, size=10)
for i, rd in enumerate(rows_data):
    for j, v in enumerate(rd):
        cell = t.rows[i+1].cells[j]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(v)
        set_font(r, size=10)
doc.add_paragraph()

heading2(doc, "8.2  Queries System")
body(doc, "The Queries module allows students to raise questions about published marks and receive responses from lecturers.")
t2 = doc.add_table(rows=6, cols=2)
t2.style = 'Table Grid'
q_data = [
    ("Feature",          "Description"),
    ("Submit Query",     "Student selects course + assessment and types a question."),
    ("Unread Badge",     "Red dot on Queries sidebar link; shows count of unanswered/unread queries."),
    ("Timestamps",       "Full date & time displayed on each query card (submission and resolution)."),
    ("Filter Tabs",      "All / Open / Resolved tabs with item counts."),
    ("Sort Order",       "Toggle between Newest first and Oldest first."),
]
for i, (a, b) in enumerate(q_data):
    t2.rows[i].cells[0].paragraphs[0].clear()
    t2.rows[i].cells[1].paragraphs[0].clear()
    ra = t2.rows[i].cells[0].paragraphs[0].add_run(a)
    rb = t2.rows[i].cells[1].paragraphs[0].add_run(b)
    set_font(ra, bold=(i==0), size=10)
    set_font(rb, size=10)
doc.add_paragraph()

heading2(doc, "8.3  Semester Timeline")
body(doc, "Coordinators and Lecturers manage academic semester timelines covering coursework deadlines.")
bullet(doc, "Fields: Academic Year, Semester, Start Date, End Date, Midterm Deadline, Grade Submission Deadline, Notes.")
bullet(doc, "Note: Finals deadlines have been removed from the system.")
bullet(doc, "Timelines appear as visual bars on a mini-calendar for at-a-glance scheduling.")

heading2(doc, "8.4  Per-Lecturer Teaching Credit Limits")
body(doc, "Each lecturer can have a personalised maximum teaching credit limit per semester.")
bullet(doc, "Set by Administrators in the user profile.")
bullet(doc, "Displayed in the lecturer selection dropdown when assigning courses.")
bullet(doc, "If no limit is set, shows as No limit — assignments proceed without credit checks.")
bullet(doc, "Coordinators and HODs can override the limit with a justification reason.")

heading2(doc, "8.5  Carry Marks Summary")
body(doc, "Available under Assessment Setup > Carry Marks Summary tab.")
bullet(doc, "Shows cumulative published carry % per student for the selected course.")
bullet(doc, "Sorted by highest carry % descending.")
bullet(doc, "Grade bands: Strong (≥70%), Satisfactory (50–69%), At Risk (<50%).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. FAQs
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "9. Frequently Asked Questions")

faqs = [
    (
        "My account is stuck on 'Pending Approval'. What should I do?",
        "Contact your system Administrator. Accounts must be manually approved. "
        "The approval status screen polls every 30 seconds so the page will update automatically once approved."
    ),
    (
        "I accidentally published a mark with the wrong score. Can I fix it?",
        "Yes. A Lecturer can click Unpublish in the column header of the Grade Sheet to revert published marks "
        "back to draft status. Then correct the score and re-publish."
    ),
    (
        "Why can I not delete an assessment?",
        "Assessments can only be deleted before the schema is locked. Once locked, the structure is frozen. "
        "Contact a Coordinator or Administrator to discuss schema changes."
    ),
    (
        "My Excel import showed errors. What do I do?",
        "The import dialog shows the exact row, student identifier, assessment column, and error message for "
        "each failed row. Common causes: wrong student email/matric number, assessment name doesn't match exactly, "
        "or score is not a valid number. Correct the file and re-import."
    ),
    (
        "The lecturer I want to assign is showing as over their credit limit. Can I still assign them?",
        "Yes. Coordinators and HODs can override credit limits. A reason field will appear — enter a justification "
        "and proceed with the assignment."
    ),
    (
        "How do I change the semester timeline?",
        "Coordinators and Lecturers can edit existing timelines from the Semester Timeline page. Click the timeline "
        "entry, modify the desired fields, and click Save."
    ),
    (
        "I can't see the Queries sidebar item. Why?",
        "The Queries link is available to Students, Lecturers, and Coordinators. HOD and Admin roles do not have "
        "this nav item by default."
    ),
    (
        "How is the Carry % calculated?",
        "Carry % = sum of (student_score / max_score × weight_percentage) for each published assessment. "
        "Draft marks are excluded. Only published marks contribute to the carry total."
    ),
]

for q, a in faqs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"Q: {q}")
    set_font(r, bold=True, size=11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.25)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(f"A: {a}")
    set_font(r2, size=11, color=GREY)

doc.add_page_break()

# ── Footer note ────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CMMS v5.0  ·  Universiti Teknologi Malaysia  ·  May 2026")
set_font(r, size=9, color=GREY)

# ── Save ───────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "CMMS_User_Manual_v5.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
