#!/usr/bin/env python3
"""
CMMS Test Case Design Document Generator
Generates a professionally formatted .docx with test suites tailored to the
Carry Mark Management System (CMMS) project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell (hex without #)."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='Calibri', size_pt=11, bold=False, color_hex='000000'):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color_hex)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)

def add_heading_custom(doc, text, level=1):
    """Add a styled heading."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, font_name='Calibri', size_pt=(18 if level==1 else (14 if level==2 else 12)), bold=True, color_hex='C90031')
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(12)
    return p

def add_paragraph_custom(doc, text, bold=False, italic=False, size_pt=11, color_hex='333333', align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name='Calibri', size_pt=size_pt, bold=bold, color_hex=color_hex)
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p

def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size_pt=11, color_hex='333333')
    return p

def add_numbered(doc, text, indent_level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size_pt=11, color_hex='333333')
    return p

# ---------------------------------------------------------------------------
# Test Case Sections
# ---------------------------------------------------------------------------

def add_test_case(doc, tc_id, title, description, preconditions, steps, expected_results):
    """Add a single test case block."""
    add_paragraph_custom(doc, f"{tc_id}: {title}", bold=True, size_pt=12, color_hex='C90031', space_after=Pt(4))
    if description:
        add_paragraph_custom(doc, description, italic=True, size_pt=11, color_hex='555555', space_after=Pt(6))
    
    # Preconditions
    add_paragraph_custom(doc, "Preconditions", bold=True, size_pt=11, color_hex='000000', space_after=Pt(3))
    for pre in preconditions:
        add_bullet(doc, pre)
    
    # Steps
    add_paragraph_custom(doc, "Steps", bold=True, size_pt=11, color_hex='000000', space_after=Pt(3))
    for i, step in enumerate(steps, 1):
        add_numbered(doc, step)
    
    # Expected Results
    add_paragraph_custom(doc, "Expected Results", bold=True, size_pt=11, color_hex='000000', space_after=Pt(3))
    for exp in expected_results:
        add_bullet(doc, exp)
    
    doc.add_paragraph()  # spacing

def add_test_suite(doc, suite_id, suite_title, test_cases):
    """Add a test suite with a border table wrapper for visual appeal."""
    add_heading_custom(doc, f"TEST SUITE {suite_id}: {suite_title}", level=2)
    
    # Summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(3.5)
    table.columns[2].width = Inches(2.0)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "TC ID"
    hdr_cells[1].text = "Test Case Title"
    hdr_cells[2].text = "Priority"
    for cell in hdr_cells:
        set_cell_shading(cell, 'C90031')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, bold=True, color_hex='FFFFFF', size_pt=11)
    
    for tc in test_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = tc['id']
        row_cells[1].text = tc['title']
        row_cells[2].text = tc.get('priority', 'High')
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size_pt=10, color_hex='333333')
    
    doc.add_paragraph()
    
    # Detailed test cases
    for tc in test_cases:
        add_test_case(doc, tc['id'], tc['title'], tc.get('description', ''),
                      tc['preconditions'], tc['steps'], tc['expected_results'])

# ---------------------------------------------------------------------------
# Document Assembly
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()
    
    # Page margins
    sections = doc.sections[0]
    sections.top_margin = Inches(1)
    sections.bottom_margin = Inches(1)
    sections.left_margin = Inches(1)
    sections.right_margin = Inches(1)
    
    # -----------------------------------------------------------------------
    # Title Page
    # -----------------------------------------------------------------------
    add_paragraph_custom(doc, "", space_after=Pt(60))
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("CARRY MARK MANAGEMENT SYSTEM (CMMS)")
    set_run_font(run, font_name='Calibri', size_pt=28, bold=True, color_hex='C90031')
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("Test Case Design Document")
    set_run_font(run2, font_name='Calibri', size_pt=20, bold=True, color_hex='333333')
    
    add_paragraph_custom(doc, "", space_after=Pt(20))
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta_p.add_run("Version 1.0\nDate: May 2026\nProject: MJIIT CMMS — Bachelor of Chemical Process Engineering")
    set_run_font(run3, font_name='Calibri', size_pt=12, bold=False, color_hex='555555')
    
    doc.add_page_break()
    
    # -----------------------------------------------------------------------
    # Table of Contents (Manual for docx stability)
    # -----------------------------------------------------------------------
    add_heading_custom(doc, "Table of Contents", level=1)
    
    toc_items = [
        ("TEST SUITE 1: Authentication & Session Management", [
            "TC01: Student Login with Valid Credentials",
            "TC02: Lecturer Login with Valid Credentials",
            "TC03: Coordinator Login with Valid Credentials",
            "TC04: HOD Login with Valid Credentials",
            "TC05: Admin Login with Valid Credentials",
            "TC06: Login with Invalid Credentials",
            "TC07: Logout Functionality",
        ]),
        ("TEST SUITE 2: User Registration & Approval Workflow", [
            "TC08: Student Self-Registration",
            "TC09: Lecturer Self-Registration (Staff Domain)",
            "TC10: Admin/Coordinator/HOD Self-Registration Blocked",
            "TC11: OTP Verification Required",
            "TC12: Admin Approves Pending Lecturer",
            "TC13: Admin Rejects Pending User",
        ]),
        ("TEST SUITE 3: Course Management", [
            "TC14: Create a New Course",
            "TC15: Edit Existing Course",
            "TC16: Archive Course with Academic Records",
            "TC17: Delete Course Without Academic Records",
            "TC18: Assign Lecturer to Course",
            "TC19: Enforce 9-Credit Workload Cap",
            "TC20: HOD Override for Workload Cap",
            "TC21: Bulk Import from Curriculum Library",
        ]),
        ("TEST SUITE 4: Enrollment & Roster Management", [
            "TC22: Upload Student Roster via Excel",
            "TC23: Invitation Token Generation for New Students",
            "TC24: Student Accepts Invitation",
            "TC25: Manual Single Student Enrollment",
            "TC26: Drop Student from Course",
        ]),
        ("TEST SUITE 5: Assessment Configuration", [
            "TC27: Create Assessment Components",
            "TC28: Update Assessment Before Marks Exist",
            "TC29: Block Assessment Update When Marks Exist",
            "TC30: Lock Assessment Schema (100% Weight Check)",
            "TC31: Delete Assessment with No Marks",
            "TC32: Block Delete Assessment with Marks",
        ]),
        ("TEST SUITE 6: Mark Entry & Grade Calculation", [
            "TC33: Lecturer Enters Raw Marks",
            "TC34: Normalised Score Auto-Calculation",
            "TC35: Publish Marks for an Assessment",
            "TC36: Block Publish Past Grade Submission Deadline",
            "TC37: Unpublish Marks (Coordinator/Admin Only)",
            "TC38: Flag a Mark for Review",
            "TC39: Student Views Own Grade & GPA",
            "TC40: At-Risk Student Detection",
        ]),
        ("TEST SUITE 7: Student Queries & Appeals", [
            "TC41: Student Raises a Query on a Mark",
            "TC42: Lecturer Responds to Query",
            "TC43: Query Resolution and Closure",
            "TC44: Query on Unpublished Mark Blocked",
        ]),
        ("TEST SUITE 8: Semester Timeline Management", [
            "TC45: Create Semester Timeline",
            "TC46: Grade Submission Deadline Enforcement",
            "TC47: Edit Active Timeline",
            "TC48: Delete Past Timeline",
        ]),
        ("TEST SUITE 9: Messaging & Notifications", [
            "TC49: Send Message to Single Recipient",
            "TC50: Message Threading (Reply to Parent)",
            "TC51: Role-Based Messaging Restrictions",
            "TC52: Rate Limiting on Message Send",
            "TC53: Mark Message as Read",
            "TC54: Unread Message Badge Count",
        ]),
        ("TEST SUITE 10: Audit Logs & Reporting", [
            "TC55: Audit Log Captures Course Creation",
            "TC56: Audit Log Captures Mark Publish",
            "TC57: Audit Log Captures Login Events",
            "TC58: HOD Views Department Analytics",
            "TC59: Coordinator Generates Grade Report",
        ]),
        ("TEST SUITE 11: Dashboard & Role-Based Access Control", [
            "TC60: Role-Based Dashboard Redirect",
            "TC61: Student Sidebar Navigation",
            "TC62: Lecturer Sidebar Navigation",
            "TC63: Coordinator Sidebar Navigation",
            "TC64: HOD Sidebar Navigation",
            "TC65: Admin Sidebar Navigation",
        ]),
    ]
    
    for suite_title, cases in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(suite_title)
        set_run_font(run, bold=True, size_pt=12, color_hex='C90031')
        for case in cases:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.3)
            run2 = p2.add_run(case)
            set_run_font(run2, size_pt=10, color_hex='333333')
    
    doc.add_page_break()
    
    # -----------------------------------------------------------------------
    # Test Suites Detail
    # -----------------------------------------------------------------------
    
    # TEST SUITE 1
    add_test_suite(doc, 1, "Authentication & Session Management", [
        {
            "id": "TC01", "title": "Student Login with Valid Credentials", "priority": "High",
            "description": "Verifies that a registered student can log in using valid email and password and is redirected to the Student Dashboard.",
            "preconditions": [
                "Student account exists in the system with approval_status='approved' and is_active=true.",
                "System is accessible at the deployed frontend URL.",
                "Student credentials: student@graduate.utm.my / password@cmsss"
            ],
            "steps": [
                "Navigate to the CMMS login page.",
                "Select 'Student' from the role selector.",
                "Enter the registered student email address.",
                "Enter the correct password.",
                "Click the 'Log In' button.",
                "Observe the post-login redirect."
            ],
            "expected_results": [
                "The login form accepts the credentials without validation errors.",
                "A JWT token is issued and stored in localStorage and as a cookie (SameSite=Strict).",
                "The user is redirected to /dashboard/student.",
                "The user object in Zustand authStore reflects the student role."
            ]
        },
        {
            "id": "TC02", "title": "Lecturer Login with Valid Credentials", "priority": "High",
            "description": "Verifies that an approved lecturer can log in and access lecturer-specific modules.",
            "preconditions": [
                "Lecturer account exists with role='lecturer', approval_status='approved', email_verified=true.",
                "Lecturer is assigned to at least one course in the current semester."
            ],
            "steps": [
                "Navigate to the CMMS login page.",
                "Select 'Lecturer' from the role selector.",
                "Enter lecturer@utm.my and the correct password.",
                "Click 'Log In'."
            ],
            "expected_results": [
                "Login succeeds and a valid JWT is returned.",
                "Redirected to /dashboard/lecturer.",
                "Sidebar displays Lecturer navigation: My Courses, Roster Upload, Smart Grid, Assessment Setup, Queries, Messages, Profile."
            ]
        },
        {
            "id": "TC03", "title": "Coordinator Login with Valid Credentials", "priority": "High",
            "description": "Verifies coordinator login and access to coordination features.",
            "preconditions": [
                "Coordinator account exists (special role assigned to a lecturer by admin).",
                "Account is approved and active."
            ],
            "steps": [
                "Select 'Lecturer' on the login page (coordinator uses lecturer login flow).",
                "Enter coordinator@utm.my and valid password.",
                "Click 'Log In'."
            ],
            "expected_results": [
                "Login succeeds.",
                "Redirected to /dashboard/coordinator.",
                "Coordinator navigation is visible: Courses, Roster Management, Assessment Config, Semester Timeline, Messages, Flagged Marks, Reports."
            ]
        },
        {
            "id": "TC04", "title": "HOD Login with Valid Credentials", "priority": "High",
            "description": "Verifies HOD login and elevated access.",
            "preconditions": [
                "HOD account exists with role='hod' (or special role flag).",
                "Account is approved and active."
            ],
            "steps": [
                "Enter hod@utm.my and valid password.",
                "Click 'Log In'."
            ],
            "expected_results": [
                "Redirected to /dashboard/hod.",
                "HOD navigation visible: Departments, Analytics, Export, Audit Log."
            ]
        },
        {
            "id": "TC05", "title": "Admin Login with Valid Credentials", "priority": "High",
            "description": "Verifies admin login and full system access.",
            "preconditions": [
                "Admin account exists: admin@utm.my / password@cmsss.",
                "Admin account is active and approved."
            ],
            "steps": [
                "Enter admin credentials and click 'Log In'."
            ],
            "expected_results": [
                "Redirected to /dashboard/admin.",
                "Admin navigation visible: Pending Approvals, Users, Roles & Permissions, Database, Settings, System Logs."
            ]
        },
        {
            "id": "TC06", "title": "Login with Invalid Credentials", "priority": "High",
            "description": "Ensures the system rejects unregistered or incorrect credentials.",
            "preconditions": [
                "The email used is not registered OR the password is incorrect.",
                "System is online."
            ],
            "steps": [
                "Navigate to login page.",
                "Enter an unregistered email (e.g., fake@graduate.utm.my).",
                "Enter any password.",
                "Click 'Log In'."
            ],
            "expected_results": [
                "An appropriate error message is displayed (e.g., 'Invalid email or password').",
                "No JWT token is issued.",
                "The user remains on the login page."
            ]
        },
        {
            "id": "TC07", "title": "Logout Functionality", "priority": "Medium",
            "description": "Verifies that logout clears session data and redirects to the public page.",
            "preconditions": [
                "User is currently logged in (any role)."
            ],
            "steps": [
                "Click the user avatar/icon at the top-right corner.",
                "Select 'Log Out' from the dropdown menu."
            ],
            "expected_results": [
                "POST /auth/logout is called.",
                "localStorage token and cookie are cleared.",
                "Zustand authStore is reset.",
                "User is redirected to the public landing page (/)."
            ]
        },
    ])
    
    # TEST SUITE 2
    add_test_suite(doc, 2, "User Registration & Approval Workflow", [
        {
            "id": "TC08", "title": "Student Self-Registration", "priority": "High",
            "description": "Verifies that a student can self-register with a valid graduate email domain.",
            "preconditions": [
                "The student email domain is @graduate.utm.my.",
                "The email is not already registered."
            ],
            "steps": [
                "Navigate to the Signup page.",
                "Select 'Student' role.",
                "Enter a valid @graduate.utm.my email, full name, matric number, and a strong password.",
                "Click 'Sign Up'."
            ],
            "expected_results": [
                "Account is created with role='student', approval_status='pending' (or auto-approved depending on policy).",
                "An audit log entry is created for the signup event.",
                "User is redirected to a pending-approval page if manual approval is required."
            ]
        },
        {
            "id": "TC09", "title": "Lecturer Self-Registration (Staff Domain)", "priority": "High",
            "description": "Verifies lecturer signup is restricted to @utm.my domain.",
            "preconditions": [
                "Lecturer uses a valid @utm.my email address."
            ],
            "steps": [
                "Select 'Lecturer' on signup.",
                "Enter a @utm.my email and valid details.",
                "Submit the form."
            ],
            "expected_results": [
                "Registration succeeds and account is created with approval_status='pending'.",
                "Audit log records the signup."
            ]
        },
        {
            "id": "TC10", "title": "Admin/Coordinator/HOD Self-Registration Blocked", "priority": "High",
            "description": "Ensures privileged roles cannot be self-registered via the public endpoint.",
            "preconditions": [
                "User attempts to register as admin, coordinator, or hod via signup form or direct API call."
            ],
            "steps": [
                "Send a POST /auth/signup request with role='admin' (or 'coordinator', 'hod').",
                "Alternatively, inspect the frontend role selector options."
            ],
            "expected_results": [
                "The API returns HTTP 400 with message: 'Only student and lecturer roles can self-register. Admin, coordinator, and HOD accounts are provisioned by an administrator.'",
                "No user record is created in the database."
            ]
        },
        {
            "id": "TC11", "title": "OTP Verification Required", "priority": "High",
            "description": "Verifies that email ownership is confirmed via OTP before full activation.",
            "preconditions": [
                "User has just signed up.",
                "SMTP/OTP service is operational."
            ],
            "steps": [
                "Complete signup form.",
                "Check email inbox for OTP code.",
                "Enter the OTP in the verification screen.",
                "Submit."
            ],
            "expected_results": [
                "OTP email is delivered within a reasonable time.",
                "Correct OTP sets email_verified=true.",
                "Incorrect OTP shows a validation error without revealing the correct code."
            ]
        },
        {
            "id": "TC12", "title": "Admin Approves Pending Lecturer", "priority": "High",
            "description": "Verifies admin can approve a pending lecturer account.",
            "preconditions": [
                "A lecturer signup exists with approval_status='pending'.",
                "Admin is logged in."
            ],
            "steps": [
                "Navigate to Pending Approvals page.",
                "Locate the pending lecturer entry.",
                "Click 'Approve'.",
                "Confirm the action in the dialog."
            ],
            "expected_results": [
                "approval_status changes to 'approved'.",
                "approved_by is set to the admin's user ID.",
                "approved_at timestamp is recorded.",
                "Audit log captures the approval action."
            ]
        },
        {
            "id": "TC13", "title": "Admin Rejects Pending User", "priority": "Medium",
            "description": "Verifies admin can reject a pending account with a reason.",
            "preconditions": [
                "A pending user exists.",
                "Admin is logged in."
            ],
            "steps": [
                "Navigate to Pending Approvals.",
                "Click 'Reject' on a pending user.",
                "Enter a rejection reason.",
                "Confirm."
            ],
            "expected_results": [
                "approval_status changes to 'rejected'.",
                "rejection_reason is stored.",
                "User cannot log in; appropriate message shown on next login attempt."
            ]
        },
    ])
    
    # TEST SUITE 3
    add_test_suite(doc, 3, "Course Management", [
        {
            "id": "TC14", "title": "Create a New Course", "priority": "High",
            "description": "Verifies that a coordinator can create a course and lecturer_id remains nullable.",
            "preconditions": [
                "Coordinator is logged in.",
                "Course code does not already exist for the same academic_year and semester."
            ],
            "steps": [
                "Navigate to Course Management > Create Course.",
                "Fill in code (e.g., KKK212), name, credits, semester, academic year.",
                "Leave Lecturer field empty.",
                "Click 'Create Course'."
            ],
            "expected_results": [
                "Course is created with lecturer_id=null.",
                "No silent auto-assignment of a random lecturer occurs.",
                "Course appears in the unassigned list for coordinator assignment.",
                "Audit log records the creation event with old_values=null and new_values containing the course details."
            ]
        },
        {
            "id": "TC15", "title": "Edit Existing Course", "priority": "High",
            "description": "Verifies course metadata can be updated by authorized roles.",
            "preconditions": [
                "A course exists and user has edit permission (coordinator/admin)."
            ],
            "steps": [
                "Open the course edit form.",
                "Change the course name and max_students.",
                "Click 'Save Changes'."
            ],
            "expected_results": [
                "Course record is updated in Supabase.",
                "Audit log stores old and new values.",
                "A success toast/notification is shown."
            ]
        },
        {
            "id": "TC16", "title": "Archive Course with Academic Records", "priority": "High",
            "description": "Verifies that courses with assessments or enrollments cannot be hard-deleted but can be archived.",
            "preconditions": [
                "Course has at least one enrollment or one assessment linked to it."
            ],
            "steps": [
                "Attempt to delete the course.",
                "Observe the system response.",
                "Use the 'Archive' action instead."
            ],
            "expected_results": [
                "Hard delete is blocked with error: 'Cannot delete: academic records exist.'",
                "Archive action succeeds; course is hidden from active lists but retained in DB.",
                "Audit log records the archive action."
            ]
        },
        {
            "id": "TC17", "title": "Delete Course Without Academic Records", "priority": "Medium",
            "description": "Verifies that a course with no linked data can be permanently deleted.",
            "preconditions": [
                "Course has zero enrollments and zero assessments."
            ],
            "steps": [
                "Click 'Delete' on the course card.",
                "Confirm deletion in the modal."
            ],
            "expected_results": [
                "Course is removed from the database.",
                "Audit log records the deletion with old_values preserved."
            ]
        },
        {
            "id": "TC18", "title": "Assign Lecturer to Course", "priority": "High",
            "description": "Verifies coordinator can assign an available lecturer.",
            "preconditions": [
                "A course exists with no lecturer assigned.",
                "Target lecturer exists and has capacity."
            ],
            "steps": [
                "Open Course Management.",
                "Click 'Assign Lecturer' on the unassigned course.",
                "Select a lecturer from the dropdown.",
                "Confirm assignment."
            ],
            "expected_results": [
                "Course.lecturer_id is updated to the selected lecturer.",
                "Lecturer's workload bar updates to reflect new total credits.",
                "Audit log records the assignment."
            ]
        },
        {
            "id": "TC19", "title": "Enforce 9-Credit Workload Cap", "priority": "High",
            "description": "Verifies the system prevents lecturer overload beyond 9 credits per semester.",
            "preconditions": [
                "Lecturer is already assigned courses totaling 9 credits in the current semester."
            ],
            "steps": [
                "Attempt to assign an additional 3-credit course to the same lecturer.",
                "Observe the system response."
            ],
            "expected_results": [
                "Assignment is blocked with message: 'Lecturer exceeds 9-credit limit for this semester.'",
                "No database update occurs."
            ]
        },
        {
            "id": "TC20", "title": "HOD Override for Workload Cap", "priority": "Medium",
            "description": "Verifies HOD can override the 9-credit cap with audit trail.",
            "preconditions": [
                "HOD is logged in.",
                "Target lecturer is at or above 9 credits."
            ],
            "steps": [
                "HOD navigates to lecturer workload view.",
                "Selects the overloaded lecturer.",
                "Clicks 'Override Cap' and provides a justification reason.",
                "Assigns the additional course."
            ],
            "expected_results": [
                "Assignment succeeds despite exceeding 9 credits.",
                "Audit log records the override with reason and HOD user ID.",
                "Workload bar shows the excess with a visual warning (e.g., red indicator)."
            ]
        },
        {
            "id": "TC21", "title": "Bulk Import from Curriculum Library", "priority": "High",
            "description": "Verifies coordinator can bulk-add canonical MJIIT curriculum courses.",
            "preconditions": [
                "Coordinator is on Course Management page.",
                "Intake year is selected (e.g., 2024)."
            ],
            "steps": [
                "Click 'Add from Curriculum Library'.",
                "Select Program: Bachelor of Chemical Process Engineering.",
                "Select Year 1, Semester 1 courses.",
                "Click 'Bulk Add to System'."
            ],
            "expected_results": [
                "All selected courses are inserted with derived academic_year based on intake year + program year offset.",
                "Courses appear in the course list with correct metadata (credits, code, name, semester).",
                "No duplicate codes are created; existing ones are skipped or warned."
            ]
        },
    ])
    
    # TEST SUITE 4
    add_test_suite(doc, 4, "Enrollment & Roster Management", [
        {
            "id": "TC22", "title": "Upload Student Roster via Excel", "priority": "High",
            "description": "Verifies coordinator can upload an Excel roster to provision multiple students.",
            "preconditions": [
                "Coordinator is logged in.",
                "An Excel file (.xlsx) is prepared with columns: email, full_name, matric_number.",
                "Active semester timeline exists."
            ],
            "steps": [
                "Navigate to Roster Management.",
                "Select the target course.",
                "Click 'Upload Roster' and choose the Excel file.",
                "Preview the parsed roster.",
                "Click 'Confirm Import'."
            ],
            "expected_results": [
                "Each new student email triggers account provisioning if not existing.",
                "Invitation tokens are generated cryptographically for new accounts.",
                "New accounts are created with is_active=false and invitation_status='pending'.",
                "Enrollment records are created with semester/academic_year derived from the active timeline.",
                "Audit log records the roster upload with number of records processed."
            ]
        },
        {
            "id": "TC23", "title": "Invitation Token Generation for New Students", "priority": "High",
            "description": "Ensures roster upload generates secure invitation tokens.",
            "preconditions": [
                "A roster upload has just been processed with new student emails."
            ],
            "steps": [
                "Inspect the users table in Supabase for the newly created student accounts.",
                "Check the invitation_token / invitation_status fields.",
                "Verify token expiration is set to 14 days from creation."
            ],
            "expected_results": [
                "Each new student has a unique, cryptographically secure invitation token.",
                "Token expiration is exactly 14 days after creation.",
                "Password hash is derived from the token (to be reset on first login).",
                "No default/weak preset password is used."
            ]
        },
        {
            "id": "TC24", "title": "Student Accepts Invitation", "priority": "High",
            "description": "Verifies a provisioned student can accept the invitation and set a real password.",
            "preconditions": [
                "Student received the invitation email containing the token/link.",
                "Token has not expired."
            ],
            "steps": [
                "Student clicks the invitation link.",
                "Enters a new strong password.",
                "Confirms the password.",
                "Submits the form."
            ],
            "expected_results": [
                "is_active becomes true.",
                "invitation_status changes to 'accepted'.",
                "Password hash is re-computed from the new password.",
                "Student can now log in with the new password."
            ]
        },
        {
            "id": "TC25", "title": "Manual Single Student Enrollment", "priority": "Medium",
            "description": "Verifies coordinator can manually add a single student to a course.",
            "preconditions": [
                "Student account already exists and is active.",
                "Coordinator is on the course roster page."
            ],
            "steps": [
                "Click 'Add Student'.",
                "Search for the student by email or matric number.",
                "Select the student.",
                "Click 'Enroll'."
            ],
            "expected_results": [
                "An enrollment record is created linking student_id and course_id.",
                "Enrollment status is set to 'enrolled'.",
                "Student appears in the course roster grid."
            ]
        },
        {
            "id": "TC26", "title": "Drop Student from Course", "priority": "Medium",
            "description": "Verifies coordinator can drop a student and the enrollment is updated.",
            "preconditions": [
                "Student is currently enrolled in the course.",
                "No published marks exist for the student in this course (or system allows soft-drop)."
            ],
            "steps": [
                "Locate the student in the roster.",
                "Click 'Drop'.",
                "Confirm the action."
            ],
            "expected_results": [
                "Enrollment status changes to 'dropped' (or record is deleted depending on policy).",
                "Student no longer appears in the active roster.",
                "Audit log records the drop action."
            ]
        },
    ])
    
    # TEST SUITE 5
    add_test_suite(doc, 5, "Assessment Configuration", [
        {
            "id": "TC27", "title": "Create Assessment Components", "priority": "High",
            "description": "Verifies lecturer/coordinator can create multiple assessment items for a course.",
            "preconditions": [
                "Course exists and lecturer is assigned.",
                "User has permission to configure assessments (lecturer of the course, coordinator, or admin)."
            ],
            "steps": [
                "Navigate to Assessment Setup.",
                "Select the course.",
                "Click 'Add Assessment'.",
                "Enter name (e.g., 'Midterm Exam'), type='exam', max_score=100, weight_percentage=30.",
                "Click 'Save'."
            ],
            "expected_results": [
                "Assessment is created and linked to the course.",
                "is_locked defaults to false.",
                "Assessment appears in the assessment list with correct weight."
            ]
        },
        {
            "id": "TC28", "title": "Update Assessment Before Marks Exist", "priority": "High",
            "description": "Verifies assessment fields can be edited when no marks are entered.",
            "preconditions": [
                "Assessment exists with zero mark records."
            ],
            "steps": [
                "Open the assessment edit form.",
                "Change max_score from 100 to 50.",
                "Update weight_percentage from 30 to 25.",
                "Save."
            ],
            "expected_results": [
                "Update succeeds.",
                "Audit log records the old and new values."
            ]
        },
        {
            "id": "TC29", "title": "Block Assessment Update When Marks Exist", "priority": "High",
            "description": "Ensures grading integrity by locking assessment schema once marks exist.",
            "preconditions": [
                "Assessment has at least one mark record entered."
            ],
            "steps": [
                "Open the assessment edit form.",
                "Attempt to change max_score or weight_percentage.",
                "Save."
            ],
            "expected_results": [
                "System returns HTTP 409 CONFLICT.",
                "Error message: 'Cannot modify assessment: marks already exist. Unpublish or delete marks first.'",
                "No database update occurs."
            ]
        },
        {
            "id": "TC30", "title": "Lock Assessment Schema (100% Weight Check)", "priority": "High",
            "description": "Verifies the lock endpoint validates total weight equals 100%.",
            "preconditions": [
                "Course has assessments with cumulative weight != 100% (e.g., 30+20=70)."
            ],
            "steps": [
                "Click 'Lock Assessment Schema'.",
                "Observe system response.",
                "Add missing assessments to reach 100%.",
                "Click 'Lock Assessment Schema' again."
            ],
            "expected_results": [
                "First attempt fails with message: 'Total weight must equal 100% before locking. Current: 70%.'",
                "After adjustment, lock succeeds.",
                "is_locked becomes true for all course assessments.",
                "Further destructive edits are blocked until explicitly unlocked."
            ]
        },
        {
            "id": "TC31", "title": "Delete Assessment with No Marks", "priority": "Medium",
            "description": "Verifies clean-up of unused assessments.",
            "preconditions": [
                "Assessment exists with zero marks."
            ],
            "steps": [
                "Click 'Delete' on the assessment.",
                "Confirm deletion."
            ],
            "expected_results": [
                "Assessment is removed.",
                "Audit log records the deletion."
            ]
        },
        {
            "id": "TC32", "title": "Block Delete Assessment with Marks", "priority": "High",
            "description": "Ensures assessments with marks cannot be accidentally deleted.",
            "preconditions": [
                "Assessment has one or more mark records."
            ],
            "steps": [
                "Click 'Delete' on the assessment.",
                "Confirm deletion."
            ],
            "expected_results": [
                "Deletion is blocked with error: 'Assessment has marks. Remove marks first or archive the course.'",
                "Assessment remains in the database."
            ]
        },
    ])
    
    # TEST SUITE 6
    add_test_suite(doc, 6, "Mark Entry & Grade Calculation", [
        {
            "id": "TC33", "title": "Lecturer Enters Raw Marks", "priority": "High",
            "description": "Verifies lecturers can enter raw scores for each student and assessment.",
            "preconditions": [
                "Lecturer is logged in.",
                "Assessment exists and is not locked (or marks are allowed).",
                "Student is enrolled in the course."
            ],
            "steps": [
                "Navigate to Smart Grid or Assessment Setup.",
                "Select the course and assessment.",
                "Enter a raw_score for a student (e.g., 75).",
                "Click 'Save'."
            ],
            "expected_results": [
                "Mark record is created with status='draft'.",
                "normalised_score is auto-calculated based on max_score.",
                "is_flagged defaults to false."
            ]
        },
        {
            "id": "TC34", "title": "Normalised Score Auto-Calculation", "priority": "High",
            "description": "Ensures normalised_score is derived correctly from raw_score and max_score.",
            "preconditions": [
                "Assessment max_score = 100.",
                "Lecturer enters raw_score = 75."
            ],
            "steps": [
                "Enter raw_score = 75 for a student.",
                "Save.",
                "Query the marks table."
            ],
            "expected_results": [
                "normalised_score = (75 / 100) * 100 = 75.0.",
                "If max_score were 50 and raw_score 40, normalised_score = 80.0."
            ]
        },
        {
            "id": "TC35", "title": "Publish Marks for an Assessment", "priority": "High",
            "description": "Verifies lecturer can publish marks so students can view them.",
            "preconditions": [
                "Draft marks exist for the assessment.",
                "Grade submission deadline has NOT passed (or override is granted)."
            ],
            "steps": [
                "Select the assessment.",
                "Click 'Publish Marks'.",
                "Confirm in the modal."
            ],
            "expected_results": [
                "All draft marks for that assessment change status to 'published'.",
                "Students can now view their marks on 'My Marks' page.",
                "Audit log records the publish event with assessment_id."
            ]
        },
        {
            "id": "TC36", "title": "Block Publish Past Grade Submission Deadline", "priority": "High",
            "description": "Ensures semester timeline deadlines are enforced.",
            "preconditions": [
                "Current date is past the grade_submission_deadline for the active semester timeline."
            ],
            "steps": [
                "Attempt to publish marks for an assessment.",
                "Observe the system response."
            ],
            "expected_results": [
                "Publish is blocked with error: 'Grade submission deadline has passed. Contact coordinator or HOD for an extension.'",
                "Marks remain in draft status."
            ]
        },
        {
            "id": "TC37", "title": "Unpublish Marks (Coordinator/Admin Only)", "priority": "High",
            "description": "Verifies only elevated roles can unpublish marks with a documented reason.",
            "preconditions": [
                "Marks are currently published.",
                "Coordinator is logged in (or Admin/HOD)."
            ],
            "steps": [
                "Navigate to the assessment marks view.",
                "Click 'Unpublish'.",
                "Enter a reason (e.g., 'Data entry error discovered during review').",
                "Confirm."
            ],
            "expected_results": [
                "Marks revert to status='draft'.",
                "Students can no longer see the marks.",
                "Audit log records the unpublish action with reason and user ID.",
                "A lecturer attempting the same action receives HTTP 403."
            ]
        },
        {
            "id": "TC38", "title": "Flag a Mark for Review", "priority": "Medium",
            "description": "Verifies lecturers/coordinators can flag suspicious or disputed marks.",
            "preconditions": [
                "A mark record exists.",
                "User has permission to flag marks (lecturer of the course, coordinator, admin)."
            ],
            "steps": [
                "Locate the mark in Smart Grid.",
                "Click 'Flag'.",
                "Enter a flag_note (e.g., 'Student claims attendance bonus not applied').",
                "Save."
            ],
            "expected_results": [
                "is_flagged becomes true.",
                "flag_note is stored.",
                "Flagged mark appears in the 'Flagged Marks' page for coordinators.",
                "An optional ai_anomaly_flag may also be set if ML-based anomaly detection is enabled."
            ]
        },
        {
            "id": "TC39", "title": "Student Views Own Grade & GPA", "priority": "High",
            "description": "Verifies students see accurate letter grades and GPA calculations based on UTM regulations.",
            "preconditions": [
                "Student is enrolled in courses with published marks.",
                "Assessments are locked and all weights sum to 100%."
            ],
            "steps": [
                "Student logs in and navigates to 'My Marks'.",
                "Selects a course.",
                "Views the grade breakdown."
            ],
            "expected_results": [
                "Weighted total percentage is displayed correctly.",
                "Letter grade is shown (e.g., A, A-, B+, etc.) based on UTM thresholds.",
                "GPA value is displayed (e.g., 3.67 for A-).",
                "If total < 50, 'At Risk' warning banner is shown."
            ]
        },
        {
            "id": "TC40", "title": "At-Risk Student Detection", "priority": "Medium",
            "description": "Ensures the system highlights students with failing or borderline grades.",
            "preconditions": [
                "A student has a weighted total < 50 OR has multiple flagged marks."
            ],
            "steps": [
                "Coordinator navigates to Reports or Flagged Marks.",
                "Views the 'At-Risk Students' list."
            ],
            "expected_results": [
                "At-risk students are listed with course, current percentage, and risk reason.",
                "Data is sortable and filterable.",
                "Coordinator can click to view detailed grade breakdown."
            ]
        },
    ])
    
    # TEST SUITE 7
    add_test_suite(doc, 7, "Student Queries & Appeals", [
        {
            "id": "TC41", "title": "Student Raises a Query on a Mark", "priority": "High",
            "description": "Verifies a student can formally question a published mark.",
            "preconditions": [
                "Student is logged in.",
                "Mark is published (status='published')."
            ],
            "steps": [
                "Navigate to 'My Marks'.",
                "Locate the assessment mark.",
                "Click 'Raise Query'.",
                "Enter query_text (e.g., 'I believe my midterm should be 82, not 72.').",
                "Submit."
            ],
            "expected_results": [
                "A course_queries record is created linked to the mark_id.",
                "Query status is pending.",
                "Lecturer receives a notification/message about the new query."
            ]
        },
        {
            "id": "TC42", "title": "Lecturer Responds to Query", "priority": "High",
            "description": "Verifies the assigned lecturer can review and reply.",
            "preconditions": [
                "A query exists with resolved_at=null.",
                "Lecturer is assigned to the course."
            ],
            "steps": [
                "Lecturer navigates to Queries page.",
                "Opens the pending query.",
                "Reviews the mark and evidence.",
                "Enters lecturer_response (e.g., 'Verified: attendance penalty applied due to 3 absences.').",
                "Clicks 'Resolve' or 'Respond'."
            ],
            "expected_results": [
                "lecturer_response is stored.",
                "resolved_at is timestamped.",
                "Student sees the response on 'My Marks' or Queries view."
            ]
        },
        {
            "id": "TC43", "title": "Query Resolution and Closure", "priority": "Medium",
            "description": "Ensures resolved queries are visually distinct and auditable.",
            "preconditions": [
                "A query has been responded to by the lecturer."
            ],
            "steps": [
                "Student views the query status.",
                "Coordinator views the query in the master list."
            ],
            "expected_results": [
                "Status shows as 'Resolved' with timestamp.",
                "Resolved queries can be filtered separately from open queries.",
                "Audit log does not necessarily need to log every query (privacy), but mark changes resulting from queries are audited."
            ]
        },
        {
            "id": "TC44", "title": "Query on Unpublished Mark Blocked", "priority": "Medium",
            "description": "Ensures students cannot query marks that are still in draft.",
            "preconditions": [
                "Mark status is 'draft'."
            ],
            "steps": [
                "Student navigates to 'My Marks'.",
                "Attempts to click 'Raise Query' on a draft mark."
            ],
            "expected_results": [
                "'Raise Query' button is disabled or hidden.",
                "If accessed via API directly, system returns 400: 'Cannot query unpublished marks.'"
            ]
        },
    ])
    
    # TEST SUITE 8
    add_test_suite(doc, 8, "Semester Timeline Management", [
        {
            "id": "TC45", "title": "Create Semester Timeline", "priority": "High",
            "description": "Verifies coordinator/admin can define academic deadlines.",
            "preconditions": [
                "Coordinator is logged in.",
                "No existing timeline for the same academic_year + semester combination."
            ],
            "steps": [
                "Navigate to Semester Timeline page.",
                "Click 'Create Timeline'.",
                "Enter academic_year='2025/2026', semester='1'.",
                "Set start_date, end_date, midterm_deadline, grade_submission_deadline, final_deadline.",
                "Save."
            ],
            "expected_results": [
                "Timeline is inserted into semester_timelines.",
                "UNIQUE constraint prevents duplicate (academic_year, semester) pairs.",
                "Visual timeline bar renders with all deadlines clearly marked.",
                "Mini-calendar on the page highlights deadline dates."
            ]
        },
        {
            "id": "TC46", "title": "Grade Submission Deadline Enforcement", "priority": "High",
            "description": "Verifies the active timeline's grade_submission_deadline blocks late mark publishes.",
            "preconditions": [
                "A semester timeline exists with grade_submission_deadline in the past."
            ],
            "steps": [
                "Lecturer attempts to publish marks.",
                "System checks the deadline via backend/utils/session.py helper."
            ],
            "expected_results": [
                "Publish action is blocked with clear error message.",
                "session helper correctly derives active semester from the timeline table."
            ]
        },
        {
            "id": "TC47", "title": "Edit Active Timeline", "priority": "Medium",
            "description": "Verifies coordinators can extend deadlines if needed.",
            "preconditions": [
                "Timeline exists and coordinator has edit permission."
            ],
            "steps": [
                "Open the timeline edit form.",
                "Extend grade_submission_deadline by one week.",
                "Save."
            ],
            "expected_results": [
                "Timeline record is updated.",
                "Audit log records the change with old and new deadline values.",
                "Lecturers can now publish until the new extended deadline."
            ]
        },
        {
            "id": "TC48", "title": "Delete Past Timeline", "priority": "Low",
            "description": "Verifies cleanup of old timelines.",
            "preconditions": [
                "Timeline is from a previous academic year and no active enrollments reference it."
            ],
            "steps": [
                "Click 'Delete' on the past timeline.",
                "Confirm."
            ],
            "expected_results": [
                "Timeline is removed.",
                "Audit log records deletion."
            ]
        },
    ])
    
    # TEST SUITE 9
    add_test_suite(doc, 9, "Messaging & Notifications", [
        {
            "id": "TC49", "title": "Send Message to Single Recipient", "priority": "High",
            "description": "Verifies users can send internal messages.",
            "preconditions": [
                "Sender is logged in.",
                "Recipient exists in the system."
            ],
            "steps": [
                "Navigate to Messages > Compose.",
                "Select recipient from dropdown.",
                "Enter subject and body.",
                "Click 'Send'."
            ],
            "expected_results": [
                "Message record is created in messages table.",
                "is_read defaults to false.",
                "Recipient sees the message in their inbox."
            ]
        },
        {
            "id": "TC50", "title": "Message Threading (Reply to Parent)", "priority": "Medium",
            "description": "Verifies messages support threading via parent_message_id.",
            "preconditions": [
                "An existing message exists between two users.",
                "User is a participant in that thread (sender or recipient)."
            ],
            "steps": [
                "Open an existing message.",
                "Click 'Reply'.",
                "Type a reply.",
                "Send."
            ],
            "expected_results": [
                "New message has parent_message_id set to the original message id.",
                "Thread is visually grouped in the UI (inbox shows conversation chain).",
                "Non-participants cannot inject messages into the thread (403 if they attempt)."
            ]
        },
        {
            "id": "TC51", "title": "Role-Based Messaging Restrictions", "priority": "High",
            "description": "Ensures messaging respects academic hierarchy (e.g., students cannot message arbitrary admins directly without escalation).",
            "preconditions": [
                "Student is logged in.",
                "System policy restricts students from messaging HOD directly."
            ],
            "steps": [
                "Student opens Compose message.",
                "Attempts to select HOD from recipient list.",
                "Alternatively, crafts a direct API request to send to HOD."
            ],
            "expected_results": [
                "HOD does not appear in the student's recipient dropdown.",
                "API returns 403: 'You are not authorized to message this recipient. Message your lecturer or coordinator instead.'",
                "Student can message their enrolled course lecturers and coordinators."
            ]
        },
        {
            "id": "TC52", "title": "Rate Limiting on Message Send", "priority": "Medium",
            "description": "Prevents spam by limiting messages per hour per user/IP.",
            "preconditions": [
                "Rate limiter is configured (e.g., 60/hour)."
            ],
            "steps": [
                "User sends 61 messages within one hour."
            ],
            "expected_results": [
                "First 60 messages succeed.",
                "61st message returns HTTP 429 Too Many Requests with retry-after header.",
                "User sees a friendly message: 'You have reached the messaging limit. Please wait before sending more messages.'"
            ]
        },
        {
            "id": "TC53", "title": "Mark Message as Read", "priority": "Medium",
            "description": "Verifies read status tracking.",
            "preconditions": [
                "A message exists with is_read=false."
            ],
            "steps": [
                "Recipient opens the message.",
                "System auto-marks as read (or user clicks 'Mark as Read')."
            ],
            "expected_results": [
                "is_read becomes true.",
                "read_at is timestamped.",
                "Unread badge count in the sidebar decreases."
            ]
        },
        {
            "id": "TC54", "title": "Unread Message Badge Count", "priority": "Low",
            "description": "Verifies the notification badge reflects unread messages accurately.",
            "preconditions": [
                "User has 3 unread messages."
            ],
            "steps": [
                "Observe the bell icon in the header.",
                "Open one message (mark as read).",
                "Observe the badge again."
            ],
            "expected_results": [
                "Badge initially shows '3'.",
                "After reading one, badge updates to '2' without a full page refresh (optimistic UI or polling)."
            ]
        },
    ])
    
    # TEST SUITE 10
    add_test_suite(doc, 10, "Audit Logs & Reporting", [
        {
            "id": "TC55", "title": "Audit Log Captures Course Creation", "priority": "High",
            "description": "Verifies every mutating action on courses is auditable.",
            "preconditions": [
                "Coordinator creates a course."
            ],
            "steps": [
                "Create a new course.",
                "Query the audit_logs table."
            ],
            "expected_results": [
                "Audit log entry exists with action='CREATE', entity_type='course', entity_id=<course_id>.",
                "user_id is the coordinator's ID.",
                "new_values JSONB contains the full course payload.",
                "ip_address and user_agent are captured."
            ]
        },
        {
            "id": "TC56", "title": "Audit Log Captures Mark Publish", "priority": "High",
            "description": "Verifies grade publication events are immutable audit records.",
            "preconditions": [
                "Lecturer publishes marks for an assessment."
            ],
            "steps": [
                "Publish marks.",
                "Check audit_logs."
            ],
            "expected_results": [
                "action='PUBLISH_MARKS', entity_type='assessment'.",
                "old_values and new_values show the status transition (draft -> published).",
                "Timestamp is exact."
            ]
        },
        {
            "id": "TC57", "title": "Audit Log Captures Login Events", "priority": "Medium",
            "description": "Ensures security-relevant events are logged.",
            "preconditions": [
                "User logs in."
            ],
            "steps": [
                "Log in as any role.",
                "Query audit_logs filtered by action='LOGIN'."
            ],
            "expected_results": [
                "Login event is recorded with user_id, ip_address, user_agent, and timestamp.",
                "Failed logins may also be logged as action='LOGIN_FAILED'."
            ]
        },
        {
            "id": "TC58", "title": "HOD Views Department Analytics", "priority": "Medium",
            "description": "Verifies HOD dashboard displays aggregate statistics.",
            "preconditions": [
                "HOD is logged in.",
                "Department has courses, lecturers, and enrolled students."
            ],
            "steps": [
                "Navigate to HOD Dashboard > Analytics.",
                "Observe the charts and KPIs."
            ],
            "expected_results": [
                "Analytics show: total courses, total students, average course GPA, lecturer workload distribution.",
                "Charts are responsive and filterable by semester.",
                "Data matches underlying Supabase tables."
            ]
        },
        {
            "id": "TC59", "title": "Coordinator Generates Grade Report", "priority": "High",
            "description": "Verifies coordinators can export/download course grade reports.",
            "preconditions": [
                "Course has published marks for all assessments."
            ],
            "steps": [
                "Coordinator navigates to Reports.",
                "Selects a course and semester.",
                "Clicks 'Generate Grade Report'.",
                "Downloads the report."
            ],
            "expected_results": [
                "Report contains: student name, matric number, assessment breakdown, weighted total, letter grade, GPA.",
                "Report format is structured (CSV or JSON available; if docx requested, formatted table).",
                "Only students with published marks are included (or draft marks if coordinator-only view)."
            ]
        },
    ])
    
    # TEST SUITE 11
    add_test_suite(doc, 11, "Dashboard & Role-Based Access Control", [
        {
            "id": "TC60", "title": "Role-Based Dashboard Redirect", "priority": "High",
            "description": "Verifies users are routed to the correct dashboard based on their highest role privilege.",
            "preconditions": [
                "User account has a valid role (student, lecturer, coordinator, hod, or admin)."
            ],
            "steps": [
                "Log in.",
                "Observe the URL after successful authentication."
            ],
            "expected_results": [
                "Student -> /dashboard/student",
                "Lecturer (without special roles) -> /dashboard/lecturer",
                "Coordinator -> /dashboard/coordinator",
                "HOD -> /dashboard/hod",
                "Admin -> /dashboard/admin"
            ]
        },
        {
            "id": "TC61", "title": "Student Sidebar Navigation", "priority": "Medium",
            "description": "Ensures students only see permitted menu items.",
            "preconditions": [
                "Student is logged in."
            ],
            "steps": [
                "Observe the left sidebar after login."
            ],
            "expected_results": [
                "Visible: Dashboard, My Courses, My Marks, Queries, Profile.",
                "Hidden: Courses, Roster Management, Assessment Config, Reports, Semester Timeline, Admin tools."
            ]
        },
        {
            "id": "TC62", "title": "Lecturer Sidebar Navigation", "priority": "Medium",
            "description": "Ensures lecturer navigation matches their responsibilities.",
            "preconditions": [
                "Lecturer is logged in (not coordinator/hod)."
            ],
            "steps": [
                "Observe sidebar."
            ],
            "expected_results": [
                "Visible: Dashboard, My Courses, Roster Upload, Smart Grid, Assessment Setup, Queries, Messages, Profile.",
                "Hidden: Course Management (full list), Flagged Marks (coordinator view), Audit Log."
            ]
        },
        {
            "id": "TC63", "title": "Coordinator Sidebar Navigation", "priority": "Medium",
            "description": "Ensures coordinator sees management tools.",
            "preconditions": [
                "Coordinator is logged in."
            ],
            "steps": [
                "Observe sidebar."
            ],
            "expected_results": [
                "Visible: Dashboard, Courses, Roster Management, Assessment Config, Semester Timeline, Messages, Flagged Marks, Reports, Profile."
            ]
        },
        {
            "id": "TC64", "title": "HOD Sidebar Navigation", "priority": "Medium",
            "description": "Ensures HOD sees oversight tools.",
            "preconditions": [
                "HOD is logged in."
            ],
            "steps": [
                "Observe sidebar."
            ],
            "expected_results": [
                "Visible: Dashboard, Departments, Analytics, Export, Audit Log, Profile."
            ]
        },
        {
            "id": "TC65", "title": "Admin Sidebar Navigation", "priority": "Medium",
            "description": "Ensures admin sees system-level controls.",
            "preconditions": [
                "Admin is logged in."
            ],
            "steps": [
                "Observe sidebar."
            ],
            "expected_results": [
                "Visible: Dashboard, Pending Approvals, Users, Roles & Permissions, Database, Settings, System Logs, Profile."
            ]
        },
    ])
    
    # -----------------------------------------------------------------------
    # Footer / Closing
    # -----------------------------------------------------------------------
    doc.add_page_break()
    add_heading_custom(doc, "Document Control", level=1)
    add_paragraph_custom(doc, "Author: CMMS Development Team", bold=True)
    add_paragraph_custom(doc, "Review Cycle: Per Sprint", bold=True)
    add_paragraph_custom(doc, "Approval: HOD / QA Lead", bold=True)
    add_paragraph_custom(doc, "This document should be updated whenever new features are introduced (e.g., new routers, UI pages, or schema changes).", italic=True, color_hex='555555')
    
    # Save
    output_path = "/Users/khobaituddinsimran/Desktop/ACTIVE WORK/CMMS/CMMS_Test_Case_Design.docx"
    doc.save(output_path)
    print(f"Document generated successfully: {output_path}")

if __name__ == "__main__":
    build_document()
